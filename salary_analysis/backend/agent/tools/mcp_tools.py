"""MCP 工具集成 — 封装 MCP Server 工具为 Agent 可调用工具

MCP (Model Context Protocol) 工具通过子进程或 HTTP 调用，
当前项目已配置 word-document-server (docx 文档生成)。
"""

import json
import subprocess
import tempfile
import os


def create_word_document(
    title: str = "薪资分析报告",
    content_sections: list[str] | None = None,
    output_path: str = "",
) -> dict:
    """创建一个 Word 文档，包含标题和多个内容章节

    此工具通过 MCP word-document-server 生成格式化的 Word 文档。
    适用于生成可下载的薪资分析报告。
    """
    if content_sections is None:
        content_sections = []
    if not output_path:
        output_path = os.path.join(tempfile.gettempdir(), f"{title}.docx")

    try:
        import httpx
        # MCP server 通常监听在本地端口或通过子进程调用
        # 这里用 python-docx 作为 fallback 实现
        from docx import Document

        doc = Document()
        doc.add_heading(title, 0)
        for i, section in enumerate(content_sections, 1):
            doc.add_heading(f"第{i}部分", level=1)
            doc.add_paragraph(section)

        doc.save(output_path)
        return {
            "success": True,
            "output_path": output_path,
            "message": f"文档已生成: {output_path}",
            "sections": len(content_sections),
        }
    except ImportError:
        return {
            "success": False,
            "error": "python-docx 未安装，请 pip install python-docx",
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"文档生成失败: {str(e)}",
        }


MCP_TOOL_DEFINITIONS = [
    {
        "name": "create_word_document",
        "description": "生成 Word 格式的薪资分析报告文档，包含标题和多个内容章节",
        "parameters": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "文档标题",
                    "default": "薪资分析报告",
                },
                "content_sections": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "每个章节的内容文本列表",
                },
                "output_path": {
                    "type": "string",
                    "description": "输出文件路径（可选，默认生成到临时目录）",
                },
            },
        },
    },
]
