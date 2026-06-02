import docx
path = r"C:\Users\Lenovo\Desktop\softwarepractice\2.基线评测报告.docx"
doc = docx.Document(path)

print("=== FINAL VERIFICATION ===")
print("\n--- Table 0: Evaluation Results ---")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
    print(f"  {cells}")

print("\n--- All content paragraphs (after Section 2) ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    # Print all paragraphs from "评测指标定义" onwards
    if i >= 35:
        print(f"  [{i}] {t[:140]}")
